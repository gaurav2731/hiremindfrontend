import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.schemas.resume import ResumeData


def generate_docx_resume(resume_data: ResumeData) -> io.BytesIO:
    """Generate a clean, professional, single-page DOCX resume from structured data."""
    doc = Document()
    
    # ─── Page setup ────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    # Heading styles
    for level in [1, 2, 3]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if level == 1:
            h_style.font.size = Pt(11)
            h_style.font.bold = True
            h_style.paragraph_format.space_before = Pt(10)
            h_style.paragraph_format.space_after = Pt(4)
            # Bottom border for section headers
            h_style.paragraph_format.border = None

    def _add_section_heading(text: str):
        """Add a section heading with a bottom border line."""
        p = doc.add_heading(text, level=1)
        # Add a thin line under the heading using a bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr",
            {},
        )
        bottom = pBdr.makeelement(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom",
            {
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val": "single",
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz": "4",
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space": "1",
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color": "CCCCCC",
            },
        )
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ─── NAME & CONTACT ────────────────────────────────────────────────
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(resume_data.name)
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.name = "Calibri"
    name_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Contact line
    contact_parts = []
    if resume_data.email:
        contact_parts.append(resume_data.email)
    if resume_data.phone:
        contact_parts.append(resume_data.phone)
    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_p.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_run.font.name = "Calibri"
        contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ─── SUMMARY ───────────────────────────────────────────────────────
    if resume_data.summary:
        _add_section_heading("Professional Summary")
        summary_p = doc.add_paragraph(resume_data.summary)
        summary_p.paragraph_format.space_after = Pt(2)

    # ─── SKILLS ────────────────────────────────────────────────────────
    if resume_data.skills or resume_data.technologies:
        _add_section_heading("Skills")
        skills_text = ""
        if resume_data.technologies:
            skills_text += ", ".join(resume_data.technologies)
        if resume_data.skills and skills_text:
            skills_text += " | " + ", ".join(resume_data.skills)
        elif resume_data.skills:
            skills_text += ", ".join(resume_data.skills)
        if skills_text:
            p = doc.add_paragraph(skills_text)
            p.paragraph_format.space_after = Pt(2)

    # ─── EXPERIENCE ────────────────────────────────────────────────────
    if resume_data.experience:
        _add_section_heading("Experience")
        for exp in resume_data.experience:
            # Job title and company on one line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            title_run = p.add_run(exp.job_title)
            title_run.bold = True
            title_run.font.size = Pt(10.5)
            company_run = p.add_run(f" — {exp.company}")
            company_run.font.size = Pt(10.5)
            
            # Date range — right aligned using tab
            if exp.start_date or exp.end_date:
                date_range = f"{exp.start_date or ''} — {exp.end_date or 'Present'}"
                date_run = p.add_run(f"    ({date_range})")
                date_run.font.size = Pt(9.5)
                date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                date_run.italic = True
            
            # Responsibilities as bullet points (LLM prompt already enforces 1-page conciseness)
            for resp in exp.responsibilities:
                bp = doc.add_paragraph(resp, style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.space_before = Pt(0)
                for run in bp.runs:
                    run.font.size = Pt(10)

    # ─── PROJECTS ──────────────────────────────────────────────────────
    if resume_data.projects:
        _add_section_heading("Projects")
        for proj in resume_data.projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            name_run = p.add_run(proj.name)
            name_run.bold = True
            name_run.font.size = Pt(10.5)
            
            if proj.technologies:
                tech_run = p.add_run(f" — {', '.join(proj.technologies)}")
                tech_run.font.size = Pt(9.5)
                tech_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            
            if proj.description:
                desc_p = doc.add_paragraph(proj.description)
                desc_p.paragraph_format.space_after = Pt(1)
                for run in desc_p.runs:
                    run.font.size = Pt(10)

    # ─── EDUCATION ─────────────────────────────────────────────────────
    if resume_data.education:
        _add_section_heading("Education")
        for edu in resume_data.education:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            deg_run = p.add_run(edu.degree)
            deg_run.bold = True
            deg_run.font.size = Pt(10.5)
            
            parts = []
            if edu.field_of_study:
                parts.append(edu.field_of_study)
            parts.append(edu.institution)
            extra = p.add_run(f" — {' — '.join(parts)}")
            extra.font.size = Pt(10)
            
            if edu.end_date:
                date_run = p.add_run(f"    ({edu.end_date})")
                date_run.font.size = Pt(9.5)
                date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                date_run.italic = True
            
            if edu.gpa:
                gpa_run = p.add_run(f"  GPA: {edu.gpa}")
                gpa_run.font.size = Pt(9.5)

    # ─── CERTIFICATIONS (compact) ──────────────────────────────────────
    if resume_data.certifications:
        _add_section_heading("Certifications")
        cert_texts = []
        for cert in resume_data.certifications:
            parts = [cert.name]
            if cert.issuer:
                parts.append(cert.issuer)
            if cert.date:
                parts.append(cert.date)
            cert_texts.append(" — ".join(parts))
        cp = doc.add_paragraph(" | ".join(cert_texts))
        cp.paragraph_format.space_after = Pt(2)
        for run in cp.runs:
            run.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream
