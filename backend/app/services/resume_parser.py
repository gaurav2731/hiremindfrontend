import io
import logging
from docx import Document
from pypdf import PdfReader
from fastapi import UploadFile

from app.schemas.resume import ResumeData
from app.services.ai_client import generate_structured

logger = logging.getLogger(__name__)


async def extract_text_from_file(file: UploadFile) -> str:
    """Extract raw text from an uploaded PDF or DOCX file."""
    content = await file.read()
    filename = file.filename.lower() if file.filename else ""
    file_size = len(content)

    logger.info(f"Processing file: {file.filename} ({file_size} bytes)")

    text = ""

    if filename.endswith(".pdf"):
        try:
            reader = PdfReader(io.BytesIO(content))
            num_pages = len(reader.pages)
            logger.info(f"PDF has {num_pages} page(s)")
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                text += page_text + "\n"
                logger.info(f"  Page {i+1}: extracted {len(page_text.strip())} chars")
        except Exception as e:
            logger.error(f"Failed to read PDF: {e}")
            raise ValueError(
                f"Could not read the PDF file. The file may be corrupted or password-protected. "
                f"Please try a different file or convert it to DOCX format."
            )

    elif filename.endswith(".docx"):
        try:
            doc = Document(io.BytesIO(content))
            para_count = len(doc.paragraphs)
            logger.info(f"DOCX has {para_count} paragraph(s)")
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            logger.error(f"Failed to read DOCX: {e}")
            raise ValueError(
                f"Could not read the DOCX file. The file may be corrupted. "
                f"Please try a different file."
            )
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")

    return text.strip()


async def parse_resume(file: UploadFile) -> ResumeData:
    """Extract text from the resume and parse it into a structured ResumeData object."""
    raw_text = await extract_text_from_file(file)

    if not raw_text:
        filename = file.filename.lower() if file.filename else ""
        if filename.endswith(".pdf"):
            raise ValueError(
                "Could not extract any text from the PDF. This usually happens when the PDF is a "
                "scanned document or image-based file with no selectable text. Try these solutions:\n"
                "1. Upload a DOCX version of your resume\n"
                "2. Use a PDF created from a Word/Google Doc (not a scan)\n"
                "3. Save the PDF using 'Print as PDF' instead of scanning"
            )
        else:
            raise ValueError(
                "Could not extract any text from the provided file. "
                "Please ensure the file contains selectable text and try again."
            )

    system_prompt = (
        "You are an expert AI resume parsing agent. "
        "Your job is to read the raw text of a candidate's resume and extract all the structured information. "
        "Do not invent or hallucinate any information. If something is not present, leave it blank."
    )

    user_prompt = f"Extract structured data from the following resume text:\n\n{raw_text}"

    logger.info(f"Calling AI to parse resume ({len(raw_text)} chars)...")
    parsed_data = await generate_structured(
        prompt=user_prompt,
        response_model=ResumeData,
        system=system_prompt
    )
    logger.info("Resume parsing completed successfully")

    return parsed_data


